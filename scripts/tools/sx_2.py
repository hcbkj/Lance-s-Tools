# -*- coding: utf-8 -*-
# @Time    : 2022/8/11 16:49
# @Author  : SYH
# @File    : sx_2.py
# @Software: PyCharm

import docx
# from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # 导入库：设置对象居中、对齐等
from docx.shared import Cm

doc = docx.Document()  # 加载模板文件
# for item in doc.sections:
#     print(item.header)
#     print(item.start_type)
# changing the page margins修改页边距
sections = doc.sections
for section in sections:
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
doc.settings.odd_and_even_pages_header_footer = True  # 启动页眉页脚奇偶页不同
偶数页眉 = doc.sections[0].even_page_header  # 对偶数页进行设置，奇数页直接对节进行页眉页脚设置即可。
偶数页眉.paragraphs[0].text = "这是一个偶数页眉"
# 偶数页眉.paragraphs[1].text = "这是一个11偶数页眉"
# print('页', list(doc.sections[0]))
doc.sections[0].different_first_page_header_footer = True  # 启动页眉页脚首页不同
首页页眉 = doc.sections[0].first_page_header
首页页眉.paragraphs[0].text = "这是首页的页眉"

print('页数', len(doc.sections))
header = doc.sections[0].header
paragraph = header.paragraphs[0]
paragraph.text = "Title of my document"
print(doc.sections[0])
print(doc.settings)
doc.save(r'D:\test.docx')
