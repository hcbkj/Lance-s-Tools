# -*- coding: utf-8 -*-
# @Time    : 2022/8/11 15:22

# 山西太原页码页数打印
import os

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # 导入库：设置对象居中、对齐等
from rpalib.log import logger
from flows.flow import ScriptDef, ListDataFlow, FileItem, StringItem, SelectItem


class sx_yemianyema(ListDataFlow):
    def __init__(self, file_path, mode, page_count):
        super().__init__()
        self.file_path = file_path
        self.mode = mode
        self.page_count = page_count

    def run(self):
        try:
            doc = Document()
            # doc = Document(r'D:\test.docx')   # 打开当前路径下的已有文档
            page = int(self.page_count)
            if page <= 0:
                logger.info('当前页数不合法!')

            doc.add_paragraph("1").alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # 右对齐

            for i in range(page - 1):
                doc.add_page_break()
                if self.mode == '单面打印':
                    doc.add_paragraph(f"{i + 2}").alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # 右对齐
                elif self.mode == '双面打印':
                    if i % 2 == 0:
                        doc.add_paragraph(f"{i + 2}").alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 左对齐
                    else:
                        doc.add_paragraph(f"{i + 2}").alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # 右对齐
                else:
                    logger.info('打印模式不正确!')

            logger.info('打印完成')
            doc.save(os.path.join(file_path, 'output.docx'))
        except Exception as e:
            logger.info('打开文件失败,请检查文件是否存在!')


export = ScriptDef(
    cls=sx_yemianyema,
    group="其他",
    title="页码页数打印",
    arguments=[
        FileItem(title="需要打印的文件路径：", name="file_path"),
        SelectItem(title="选择打印模式", name="mode", options=[
            '单面打印',
            '双面打印'
        ]),
        StringItem(title="打印页数", name="page_count"),

    ]
)

if __name__ == '__main__':
    file_path = r'D:\工作相关\测试文件夹'
    mode = '双面打印'
    page_count = 10
    sx_yemianyema(file_path, mode, page_count).run()
